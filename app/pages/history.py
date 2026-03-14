from __future__ import annotations

import sys
from io import BytesIO
from pathlib import Path
from xml.sax.saxutils import escape

import streamlit as st

# Ensure project root is on path for backend imports
_ROOT = Path(__file__).resolve().parent.parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    PDF_AVAILABLE = True
except ImportError:
    letter = None
    getSampleStyleSheet = None
    Paragraph = None
    SimpleDocTemplate = None
    Spacer = None
    PDF_AVAILABLE = False


ASSETS_DIR = Path(__file__).resolve().parent.parent.parent / "assets"
CSS_FILE = ASSETS_DIR / "styles.css"


def inject_styles() -> None:
    if CSS_FILE.exists():
        st.markdown(f"<style>{CSS_FILE.read_text(encoding='utf-8')}</style>", unsafe_allow_html=True)


def clear_search() -> None:
    st.session_state["history_search_input"] = ""
    st.session_state["selected_case_id"] = None


def _to_time_text(value: object) -> str:
    if value is None:
        return ""
    if hasattr(value, "isoformat") and callable(getattr(value, "isoformat")):
        try:
            return str(value.isoformat())
        except Exception:  # noqa: BLE001
            pass
    return str(value)


def _make_json_safe(obj: object):
    try:
        import numpy as np  # type: ignore
    except Exception:  # noqa: BLE001
        np = None  # type: ignore

    try:
        import pandas as pd  # type: ignore
    except Exception:  # noqa: BLE001
        pd = None  # type: ignore

    if obj is None or isinstance(obj, (str, int, float, bool)):
        return obj

    if np is not None:
        if isinstance(obj, np.integer):
            return int(obj)
        if isinstance(obj, np.floating):
            return float(obj)
        if isinstance(obj, np.bool_):
            return bool(obj)

    if pd is not None:
        if isinstance(obj, pd.Timestamp):
            return obj.isoformat()
        if isinstance(obj, pd.Series):
            return _make_json_safe(obj.to_dict())
        if isinstance(obj, pd.DataFrame):
            return _make_json_safe(obj.to_dict(orient="records"))

    if isinstance(obj, Path):
        return str(obj)

    if isinstance(obj, dict):
        return {str(k): _make_json_safe(v) for k, v in obj.items()}

    if isinstance(obj, (list, tuple, set)):
        return [_make_json_safe(v) for v in obj]

    if hasattr(obj, "item") and callable(getattr(obj, "item")):
        try:
            item_value = obj.item()
            if item_value is not obj:
                return _make_json_safe(item_value)
        except Exception:  # noqa: BLE001
            pass

    return str(obj)


def _record_to_ui_case(rec: dict) -> dict:
    """Map Firestore history record to existing UI case format."""
    outputs = (
        rec.get("outputs")
        or rec.get("fused")
        or rec.get("fused_output")
        or rec.get("fusion_output")
        or {}
    )
    final = outputs.get("final", {}) if isinstance(outputs, dict) else {}
    motive_obj = final.get("motive", {})
    motive = motive_obj.get("pred") if isinstance(motive_obj, dict) else motive_obj

    structured = rec.get("inputs") or rec.get("structured_input") or {}
    if not isinstance(structured, dict):
        structured = {}
    created_raw = rec.get("created_at_local") or rec.get("created_at") or rec.get("timestamp")
    created_text = _to_time_text(created_raw)
    created_display = created_text or "-"

    similar_cases = rec.get("rag_results") or rec.get("similar_cases") or []
    if not isinstance(similar_cases, list):
        similar_cases = []

    risk = (
        final.get("risk_level")
        or final.get("risk_category")
        or outputs.get("final_risk_category", "-")
    )

    feedback = rec.get("feedback") or {}
    if not isinstance(feedback, dict):
        feedback = {}
    feedback_text = str(feedback.get("text") or "").strip()
    helpfulness = feedback.get("helpful") or feedback.get("helpfulness")
    notes_parts = []
    if feedback_text:
        notes_parts.append(f"Feedback: {feedback_text}")
    if helpfulness:
        notes_parts.append(f"Helpful: {helpfulness}")

    doc_id = str(rec.get("id") or rec.get("record_id") or "").strip()
    case_id = str(rec.get("case_id") or rec.get("id") or rec.get("record_id") or "").strip()
    crime_title = structured.get("primary_type") or structured.get("crime_type") or "Unknown Crime Type"
    crime_title = str(crime_title).strip()
    if crime_title.lower() in {"", "unknown", "none", "-", "nan"}:
        crime_title = "Unknown Crime Type"

    def _clean_text(value: object) -> str:
        if value is None:
            return ""
        text = str(value).strip()
        return text

    description_text = (
        _clean_text(rec.get("narrative_text"))
        or _clean_text(structured.get("narrative_text"))
        or _clean_text(structured.get("description"))
        or "No description provided."
    )
    description_preview = (
        f"{description_text[:300]}..."
        if len(description_text) > 300
        else description_text
    )

    motive_band = motive_obj.get("band") if isinstance(motive_obj, dict) else ""
    motive_conf = motive_obj.get("conf") if isinstance(motive_obj, dict) else ""
    motive_line = ""
    if motive:
        motive_line = f"Motive: {motive}"
        extra = []
        if motive_band:
            extra.append(f"Band={motive_band}")
        if motive_conf not in ("", None):
            extra.append(f"Confidence={motive_conf}")
        if extra:
            motive_line = f"{motive_line} ({', '.join(extra)})"
        notes_parts.insert(0, motive_line)

    if similar_cases:
        notes_parts.append(f"Similar cases saved: {len(similar_cases)}")
    else:
        notes_parts.append("No similar cases saved.")

    return {
        "_selector_id": doc_id or case_id,
        "doc_id": doc_id,
        "id": str(case_id),
        "title": crime_title,
        "date": created_display,
        "victim": (
            structured.get("victim_name")
            or structured.get("victim_type")
            or structured.get("victim")
            or "-"
        ),
        "age": structured.get("victim_age") or structured.get("age") or structured.get("vict_age"),
        "gender": (
            structured.get("victim_gender")
            or structured.get("victim_sex")
            or structured.get("gender")
            or "-"
        ),
        "location": (
            structured.get("location_desc")
            or structured.get("location")
            or structured.get("place")
            or structured.get("area")
            or "-"
        ),
        "description": description_text,
        "description_preview": description_preview,
        "risk": str(risk),
        "notes": " | ".join(notes_parts) if notes_parts else "",
        "motive": motive_line or "-",
        "feedback_text": feedback_text or "-",
        "helpful": str(helpfulness) if helpfulness not in (None, "") else "-",
        "inputs_full": _make_json_safe(structured),
        "outputs_full": _make_json_safe(outputs),
        "similar_cases_full": _make_json_safe(similar_cases),
        "feedback_full": _make_json_safe(feedback),
        "created_at_local": created_display,
        "_raw": rec,
        "_raw_safe": _make_json_safe(rec),
    }


def _load_cases():
    """Load full case records from Firestore history collection."""
    try:
        from services.firebase_service import list_history_records

        records = list_history_records(limit=50)
        if records:
            return [_record_to_ui_case(record) for record in records]
        return []
    except Exception:
        return []


def seed_cases():
    """Backend-connected case loader."""
    return _load_cases()


def _format_case_label(case: dict) -> str:
    return f"Case ID: {case.get('id', '-')} | {case.get('title', 'Unknown')} | {case.get('date', '-')}"


def _humanize_key(key: object) -> str:
    return str(key).replace("_", " ").strip().title()


def _format_display_value(value: object) -> str:
    if value in (None, ""):
        return "-"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, float):
        return f"{value:.3f}".rstrip("0").rstrip(".")
    if isinstance(value, dict):
        pred = value.get("pred") if isinstance(value, dict) else None
        band = value.get("band") if isinstance(value, dict) else None
        conf = value.get("conf") if isinstance(value, dict) else None
        if pred not in (None, ""):
            details = []
            if band not in (None, ""):
                details.append(f"Band: {band}")
            if conf not in (None, ""):
                details.append(f"Confidence: {_format_display_value(conf)}")
            if details:
                return f"{pred} ({', '.join(details)})"
            return str(pred)

        parts = []
        for sub_key, sub_value in value.items():
            formatted = _format_display_value(sub_value)
            if formatted == "-":
                continue
            parts.append(f"{_humanize_key(sub_key)}: {formatted}")
        return " | ".join(parts) if parts else "-"
    if isinstance(value, (list, tuple, set)):
        items = []
        for item in value:
            formatted = _format_display_value(item)
            if formatted != "-":
                items.append(formatted)
        if not items:
            return "-"
        preview = items[:5]
        suffix = "..." if len(items) > 5 else ""
        return f"{', '.join(preview)}{suffix}"
    return str(value)


def _render_compact_section(title: str, data: object, *, exclude_keys: set[str] | None = None) -> None:
    st.markdown(f"### {title}")
    if not isinstance(data, dict) or not data:
        st.write("-")
        return

    exclude = exclude_keys or set()
    rows = []
    for key, value in data.items():
        if key in exclude:
            continue
        formatted = _format_display_value(value)
        if formatted == "-":
            continue
        rows.append({"Field": _humanize_key(key), "Value": formatted})

    if not rows:
        st.write("-")
        return

    st.table(
        {
            "Field": [row["Field"] for row in rows],
            "Value": [row["Value"] for row in rows],
        }
    )


def _render_rag_evidence(case: dict) -> None:
    st.markdown("### RAG evidence")
    rag_items = case.get("similar_cases_full") or []
    if not isinstance(rag_items, list) or not rag_items:
        st.write("No RAG evidence saved.")
        return

    try:
        import pandas as pd

        df_rag = pd.DataFrame(rag_items)
        preferred = [
            "score",
            "case_id",
            "type",
            "place",
            "area",
            "risk_level",
            "motive",
            "victim_age",
            "victim_sex",
            "suspect_age",
            "suspect_sex",
        ]
        cols = [c for c in preferred if c in df_rag.columns]
        if not cols:
            cols = list(df_rag.columns)
        st.dataframe(df_rag.reindex(columns=cols), use_container_width=True, hide_index=True)
    except Exception:
        st.write("Saved RAG evidence could not be rendered as a table.")


def _render_export_button(case: dict) -> None:
    safe_case_id = "".join(
        char if char.isalnum() or char in {"-", "_"} else "_"
        for char in str(case.get("id") or "case")
    )
    if DOCX_AVAILABLE:
        st.download_button(
            "Export DOCX",
            data=build_case_docx(case),
            file_name=f"{safe_case_id}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        return

    if PDF_AVAILABLE:
        st.download_button(
            "Export PDF",
            data=build_case_pdf(case),
            file_name=f"{safe_case_id}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
        return

    st.download_button(
        "Export RTF",
        data=build_case_rtf(case),
        file_name=f"{safe_case_id}.rtf",
        mime="application/rtf",
        use_container_width=True,
    )


def build_case_docx(case: dict) -> bytes:
    """Create a simple Word document for a case export."""
    if not DOCX_AVAILABLE or Document is None:
        raise RuntimeError("python-docx is not installed")
    doc = Document()
    doc.add_heading(case["title"], level=1)

    doc.add_paragraph(f"Case ID: {case['id']}")
    doc.add_paragraph(f"Date: {case['date']}")
    doc.add_paragraph(f"Location: {case['location']}")
    doc.add_paragraph(f"Risk: {case['risk']}")

    doc.add_heading("Victim / profile info", level=2)
    doc.add_paragraph(f"Victim: {case['victim']}")
    age_value = case.get("age", "--") if case.get("age") is not None else "--"
    doc.add_paragraph(f"Age: {age_value}")
    doc.add_paragraph(f"Gender: {case.get('gender', '--')}")

    doc.add_heading("Description", level=2)
    doc.add_paragraph(case["description"])

    doc.add_heading("Notes / recommendations", level=2)
    doc.add_paragraph(case["notes"])

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_case_rtf(case: dict) -> bytes:
    """Create a Word-compatible RTF export without external dependencies."""

    def safe(value: object) -> str:
        text = str(value) if value is not None else ""
        text = text.replace("\\", r"\\").replace("{", r"\{").replace("}", r"\}")
        text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\n", r"\line ")
        return text

    age_value = case.get("age", "--") if case.get("age") is not None else "--"
    sections = [
        (r"\b " + safe(case.get("title", "Case Export")) + r"\b0", ""),
        ("Case ID:", safe(case.get("id", "--"))),
        ("Date:", safe(case.get("date", "--"))),
        ("Location:", safe(case.get("location", "--"))),
        ("Risk:", safe(case.get("risk", "--"))),
        (r"\b Victim / profile info\b0", ""),
        ("Victim:", safe(case.get("victim", "--"))),
        ("Age:", safe(age_value)),
        ("Gender:", safe(case.get("gender", "--"))),
        (r"\b Description\b0", ""),
        ("", safe(case.get("description", ""))),
        (r"\b Notes / recommendations\b0", ""),
        ("", safe(case.get("notes", ""))),
    ]

    lines = [r"{\rtf1\ansi\deff0"]
    for label, value in sections:
        if label and value:
            lines.append(rf"{label} {value}\line ")
        elif label:
            lines.append(rf"{label}\line ")
        else:
            lines.append(rf"{value}\line ")
    lines.append("}")
    return "".join(lines).encode("utf-8")


def build_case_pdf(case: dict) -> bytes:
    """Create a simple PDF document for a case export."""
    if not PDF_AVAILABLE or SimpleDocTemplate is None:
        raise RuntimeError("reportlab is not installed")

    def safe(value: object) -> str:
        return escape(str(value)) if value is not None else ""

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, title=str(case.get("title", "Case Export")))
    styles = getSampleStyleSheet()
    story = [
        Paragraph(safe(case.get("title", "Case Export")), styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"<b>Case ID:</b> {safe(case.get('id', '--'))}", styles["BodyText"]),
        Paragraph(f"<b>Date:</b> {safe(case.get('date', '--'))}", styles["BodyText"]),
        Paragraph(f"<b>Location:</b> {safe(case.get('location', '--'))}", styles["BodyText"]),
        Paragraph(f"<b>Risk:</b> {safe(case.get('risk', '--'))}", styles["BodyText"]),
        Spacer(1, 10),
        Paragraph("Victim / profile info", styles["Heading2"]),
        Paragraph(f"<b>Victim:</b> {safe(case.get('victim', '--'))}", styles["BodyText"]),
        Paragraph(
            f"<b>Age:</b> {safe(case.get('age', '--') if case.get('age') is not None else '--')}",
            styles["BodyText"],
        ),
        Paragraph(f"<b>Gender:</b> {safe(case.get('gender', '--'))}", styles["BodyText"]),
        Spacer(1, 10),
        Paragraph("Description", styles["Heading2"]),
        Paragraph(safe(case.get("description", "")), styles["BodyText"]),
        Spacer(1, 10),
        Paragraph("Notes / recommendations", styles["Heading2"]),
        Paragraph(safe(case.get("notes", "")), styles["BodyText"]),
    ]
    doc.build(story)
    return buffer.getvalue()


def render_case_details(case: dict, show_title: bool = True) -> None:
    c = {k: v for k, v in case.items() if k != "_raw"}
    if show_title:
        st.subheader(c["title"])
        st.caption(f"Case ID: {c.get('id', '-')}")
    meta_cols = st.columns(3)
    meta_cols[0].markdown(f"**Date:** {c['date']}")
    meta_cols[1].markdown(f"**Location:** {c['location']}")
    meta_cols[2].markdown(f"**Risk:** {c['risk']}")

    st.markdown("### Victim / profile info")
    info_cols = st.columns(3)
    info_cols[0].markdown(f"**Victim:** {c['victim']}")
    info_cols[1].markdown(f"**Age:** {c.get('age','--') if c.get('age') is not None else '--'}")
    info_cols[2].markdown(f"**Gender:** {c.get('gender','--')}")

    _render_compact_section(
        "Inputs",
        c.get("inputs_full"),
        exclude_keys={"narrative_text", "description", "victim_name", "victim_age", "victim_gender", "victim_sex"},
    )

    st.markdown("### Description")
    st.write(c["description"])

    output_payload = c.get("outputs_full")
    if isinstance(output_payload, dict) and isinstance(output_payload.get("final"), dict):
        compact_outputs = dict(output_payload.get("final") or {})
        if "risk_level" not in compact_outputs and output_payload.get("final_risk_category") not in (None, ""):
            compact_outputs["risk_level"] = output_payload.get("final_risk_category")
        if output_payload.get("combined_confidence") not in (None, ""):
            compact_outputs["combined_confidence"] = output_payload.get("combined_confidence")
        if output_payload.get("model_version") not in (None, ""):
            compact_outputs["model_version"] = output_payload.get("model_version")
    else:
        compact_outputs = output_payload

    _render_compact_section(
        "Outputs",
        compact_outputs,
        exclude_keys={"explanations", "fusion_meta", "tabular", "nlp", "warnings"},
    )

    st.markdown("### Notes / recommendations")
    st.markdown(c.get("motive", "-"))
    st.markdown(f"**Feedback:** {c.get('feedback_text', '-')}")
    st.markdown(f"**Helpful:** {c.get('helpful', '-')}")

    _render_rag_evidence(c)


def main() -> None:
    st.set_page_config(
        page_title="Crime Sense | Past Predictions",
        page_icon="P",
        layout="wide",
        initial_sidebar_state="collapsed",
    )
    inject_styles()

    if "history_search_input" not in st.session_state:
        st.session_state["history_search_input"] = ""
    if "selected_case_id" not in st.session_state:
        st.session_state["selected_case_id"] = None

    header_left, header_right = st.columns([3, 1.2])
    with header_left:
        st.title("Past Predicted Profiles")
        st.caption("Review prior predictions; select a case to view its details.")
    with header_right:
        s_col, x_col = st.columns([0.82, 0.18], gap="small")
        with s_col:
            st.text_input(
                "Search cases",
                placeholder="Search by case ID or crime type",
                label_visibility="collapsed",
                key="history_search_input",
            )
        with x_col:
            st.button("X", help="Clear search and show all cases", on_click=clear_search, key="clear_history_search")

    cases = seed_cases()
    query = st.session_state.get("history_search_input", "").strip().lower()
    if query:
        cases = [
            c
            for c in cases
            if query in c["id"].lower()
            or query in c["title"].lower()
        ]
    cases_sorted = sorted(cases, key=lambda c: c["date"], reverse=True)

    selected_case = None
    if cases_sorted:
        selectable_ids = [c.get("_selector_id") for c in cases_sorted if c.get("_selector_id")]
        if st.session_state.get("selected_case_id") not in selectable_ids:
            st.session_state["selected_case_id"] = selectable_ids[0] if selectable_ids else None
        selected_selector_id = st.session_state.get("selected_case_id")
        selected_case = next(
            (c for c in cases_sorted if c.get("_selector_id") == selected_selector_id),
            None,
        )
    else:
        st.session_state["selected_case_id"] = None

    list_col, detail_col = st.columns([1.2, 2.3], gap="large")

    with list_col:
        st.markdown("**Cases**")
        if not cases_sorted:
            st.info("No cases match your search.")
        else:
            st.radio(
                "Saved cases",
                options=[case.get("_selector_id") for case in cases_sorted if case.get("_selector_id")],
                format_func=lambda selector_id: _format_case_label(
                    next(case for case in cases_sorted if case.get("_selector_id") == selector_id)
                ),
                key="selected_case_id",
                label_visibility="collapsed",
            )
            selected_case = next(
                (
                    c
                    for c in cases_sorted
                    if c.get("_selector_id") == st.session_state.get("selected_case_id")
                ),
                selected_case,
            )

    with detail_col:
        if selected_case:
            action_col, _ = st.columns([1, 2.2])
            with action_col:
                _render_export_button(selected_case)
            render_case_details(selected_case)
        elif cases:
            st.info("Select a case from the left to view its details.")
        else:
            st.info("No saved cases are available.")


if __name__ == "__main__":
    main()
