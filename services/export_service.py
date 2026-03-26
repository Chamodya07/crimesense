from __future__ import annotations

from io import BytesIO
from typing import Any, TYPE_CHECKING

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument


def safe_text(value: Any) -> str:
    if value is None:
        return ""

    try:
        import numpy as np  # type: ignore
    except Exception:  # noqa: BLE001
        np = None  # type: ignore

    try:
        import pandas as pd  # type: ignore
    except Exception:  # noqa: BLE001
        pd = None  # type: ignore

    if np is not None and isinstance(value, (np.integer, np.floating, np.bool_)):
        value = value.item()

    if pd is not None and isinstance(value, pd.Timestamp):
        return value.isoformat()

    if isinstance(value, bool):
        return "Yes" if value else "No"

    if isinstance(value, float):
        if value.is_integer():
            return str(int(value))
        return str(value)

    if isinstance(value, (list, tuple, set)):
        return ", ".join(part for part in (safe_text(item) for item in value) if part)

    if isinstance(value, dict):
        parts = []
        for key, item in value.items():
            text = safe_text(item)
            if text:
                parts.append(f"{key}: {text}")
        return " | ".join(parts)

    return str(value)


def add_table(doc: "DocxDocument", rows: list[dict], cols: list[str], max_rows: int = 10) -> None:
    usable_rows = [row for row in rows if isinstance(row, dict)][:max_rows]
    usable_cols = [col for col in cols if any(safe_text(row.get(col)) for row in usable_rows)] or cols
    if not usable_rows or not usable_cols:
        return

    table = doc.add_table(rows=1, cols=len(usable_cols))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, col in enumerate(usable_cols):
        header_cells[idx].text = str(col).replace("_", " ").title()

    for row in usable_rows:
        cells = table.add_row().cells
        for idx, col in enumerate(usable_cols):
            cells[idx].text = safe_text(row.get(col))


def _add_field(doc: "DocxDocument", label: str, value: Any) -> None:
    text = safe_text(value)
    if text:
        doc.add_paragraph(f"{label}: {text}")


def build_profile_doc(record: dict) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed") from exc

    doc = Document()
    doc.add_heading("CrimeSense - Case Profile Report", level=0)

    inputs = record.get("inputs") if isinstance(record.get("inputs"), dict) else {}
    outputs = record.get("outputs") if isinstance(record.get("outputs"), dict) else {}
    final = outputs.get("final") if isinstance(outputs.get("final"), dict) else {}
    feedback = record.get("feedback") if isinstance(record.get("feedback"), dict) else {}

    rag_payload = record.get("rag") if isinstance(record.get("rag"), dict) else {}
    rag_summary = record.get("rag_summary") if isinstance(record.get("rag_summary"), dict) else {}
    if not rag_summary:
        rag_summary = rag_payload.get("summary") if isinstance(rag_payload.get("summary"), dict) else {}

    rag_results = record.get("rag_results")
    if not isinstance(rag_results, list):
        rag_results = rag_payload.get("top_cases") if isinstance(rag_payload.get("top_cases"), list) else []

    similar_saved_cases = record.get("similar_saved_cases")
    if not isinstance(similar_saved_cases, list):
        similar_saved_cases = (
            rag_payload.get("similar_saved_cases") if isinstance(rag_payload.get("similar_saved_cases"), list) else []
        )

    doc.add_heading("Case Details", level=1)
    _add_field(doc, "Case ID", record.get("case_id") or record.get("id"))
    _add_field(doc, "Timestamp", record.get("timestamp") or record.get("created_at_local") or record.get("date"))
    _add_field(doc, "City", inputs.get("city") or record.get("city"))
    _add_field(doc, "Primary Type", inputs.get("primary_type") or inputs.get("type") or record.get("title"))
    _add_field(doc, "Location Desc", inputs.get("location_desc") or inputs.get("place") or record.get("location"))
    _add_field(doc, "Weapon Desc", inputs.get("weapon_desc") or inputs.get("weapon"))

    victim = inputs.get("victim") if isinstance(inputs.get("victim"), dict) else {}
    doc.add_heading("Victim Info", level=1)
    _add_field(doc, "Victim Age", victim.get("age") or inputs.get("victim_age") or record.get("age"))
    _add_field(doc, "Victim Sex", victim.get("gender") or inputs.get("victim_sex") or record.get("gender"))
    _add_field(doc, "Victim Race", victim.get("race") or inputs.get("victim_race"))

    narrative_text = (
        record.get("narrative_text")
        or inputs.get("narrative_text")
        or record.get("description")
        or ""
    )
    if safe_text(narrative_text):
        doc.add_heading("Narrative / Description", level=1)
        doc.add_paragraph(safe_text(narrative_text))

    doc.add_heading("Predictions", level=1)
    _add_field(doc, "Risk Level", final.get("risk_level") or final.get("risk_category") or record.get("risk"))
    _add_field(doc, "Crime Severity", final.get("crime_severity"))
    _add_field(doc, "Offender Experience", final.get("offender_experience"))
    motive = final.get("motive")
    if isinstance(motive, dict):
        _add_field(doc, "Motive", motive.get("pred"))
    else:
        _add_field(doc, "Motive", motive)

    if rag_summary:
        doc.add_heading("RAG Evidence Summary", level=1)
        for key, value in rag_summary.items():
            text = safe_text(value)
            if text:
                doc.add_paragraph(f"{str(key).replace('_', ' ').title()}: {text}", style="List Bullet")

    if rag_results:
        doc.add_heading("Top Similar Cases", level=1)
        rag_cols = [
            "score",
            "case_id",
            "type",
            "place",
            "area",
            "victim_age",
            "victim_sex",
            "victim_race",
            "suspect_age",
            "suspect_sex",
            "suspect_race",
        ]
        add_table(doc, rag_results, rag_cols, max_rows=10)

    if similar_saved_cases:
        doc.add_heading("Similar Saved Cases", level=1)
        saved_cols = [
            "score",
            "saved_time",
            "record_id",
            "type",
            "weapon",
            "place",
            "city",
            "area",
            "victim_age",
            "victim_gender",
            "victim_race",
            "suspect_age",
            "suspect_sex",
            "suspect_race",
        ]
        add_table(doc, similar_saved_cases, saved_cols, max_rows=10)

    if feedback:
        doc.add_heading("Feedback", level=1)
        _add_field(doc, "Helpful", feedback.get("helpful"))
        _add_field(doc, "Feedback Text", feedback.get("text"))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


__all__ = ["build_profile_doc", "safe_text", "add_table"]
